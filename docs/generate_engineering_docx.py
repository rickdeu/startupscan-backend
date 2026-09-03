import os
from datetime import datetime

from docx import Document
from docx.shared import Inches


ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
DOCS_DIR = os.path.join(ROOT_DIR, "docs")
ASSETS_DIR = os.path.join(DOCS_DIR, "assets")
DOCX_PATH = os.path.join(DOCS_DIR, "Software_Engineering_Documentation.docx")


def _add_bullet_list(document: Document, items: list[str]):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def _add_numbered_list(document: Document, items: list[str]):
    for item in items:
        document.add_paragraph(item, style="List Number")


def _add_table(document: Document, headers: list[str], rows: list[list[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    return table


def build_docx():
    os.makedirs(DOCS_DIR, exist_ok=True)

    doc = Document()
    doc.add_heading("Software Engineering Documentation", level=0)
    doc.add_paragraph("StartupScan - complete and detailed version")
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    doc.add_heading("1. Summary and objectives", level=1)
    doc.add_paragraph(
        "StartupScan is an AI-powered platform for evaluating startups, designed to reduce analysis time, "
        "increase feedback standardization, and produce executive artifacts for decision-making."
    )
    _add_bullet_list(
        doc,
        [
            "Receive pitches in multimodal format.",
            "Generate an interpretable score and recommendations.",
            "Deliver a technical report, AI video, and visual pitch deck.",
            "Support continuous operation with job monitoring.",
        ],
    )

    doc.add_heading("2. Implemented functional scope", level=1)
    _add_bullet_list(
        doc,
        [
            "Multimodal submission (text, document, audio, video, and YouTube).",
            "Startup analysis with a 0-10 score and category-based recommendations.",
            "Local engine and GPT option with fallback.",
            "AI video in the result (D-ID/local/hybrid) with real-time progress.",
            "Slide-style pitch PDF with automatic context-based design.",
            "Pitch PDF with manual premium mode and user-selected template.",
            "Model management with training, retraining, activation, and real-time monitoring.",
            "Operational and investor dashboards.",
        ],
    )

    doc.add_heading("3. Technical architecture", level=1)
    doc.add_paragraph(
        "The solution uses a modular architecture with specialized services for analysis, video, and export."
    )
    _add_table(
        doc,
        ["Layer", "Technology", "Responsibility"],
        [
            ["Frontend", "Django Templates + JS", "Interface, forms, dashboards, and progress polling"],
            ["Backend", "Django + DRF", "Orchestration, business rules, routes, and security"],
            ["AI", "scikit-learn + OpenAI", "Scoring, interpretability, and narrative"],
            ["Video", "moviepy + D-ID", "AI video generation and local rendering"],
            ["Documentation", "reportlab + python-docx", "PDF and DOCX export"],
            ["Persistence", "SQLite/PostgreSQL", "Analyses, submissions, and operational metadata"],
        ],
    )

    architecture_image = os.path.join(ASSETS_DIR, "platform_architecture.png")
    if os.path.exists(architecture_image):
        doc.add_paragraph("Architecture diagram:")
        doc.add_picture(architecture_image, width=Inches(6.4))

    flow_image = os.path.join(ASSETS_DIR, "functional_flow.png")
    if os.path.exists(flow_image):
        doc.add_paragraph("Functional flow:")
        doc.add_picture(flow_image, width=Inches(6.4))

    categories_image = os.path.join(ASSETS_DIR, "category_example.png")
    if os.path.exists(categories_image):
        doc.add_paragraph("Example of evaluation categories:")
        doc.add_picture(categories_image, width=Inches(6.4))

    jobs_image = os.path.join(ASSETS_DIR, "job_phases.png")
    if os.path.exists(jobs_image):
        doc.add_paragraph("Asynchronous job phases:")
        doc.add_picture(jobs_image, width=Inches(6.4))

    doc.add_heading("4. Business flows", level=1)
    doc.add_paragraph("Flow A - Multimodal evaluation:")
    _add_numbered_list(
        doc,
        [
            "Receive multimodal input.",
            "Extract and consolidate context.",
            "Run local/GPT inference.",
            "Persist the structured result.",
            "Display score and recommendations.",
        ],
    )
    doc.add_paragraph("Flow B - AI video:")
    _add_numbered_list(
        doc,
        [
            "Select mode (auto/did_only/local_only).",
            "Create an asynchronous job.",
            "Track progress via endpoint.",
            "Persist artifact and metadata.",
        ],
    )
    doc.add_paragraph("Flow C - Pitch deck PDF:")
    _add_numbered_list(
        doc,
        [
            "Build the pitch narrative.",
            "Select the design mode (automatic/manual).",
            "Render visual slides.",
            "Deliver the PDF for download.",
        ],
    )

    doc.add_heading("5. Main endpoints", level=1)
    _add_table(
        doc,
        ["Endpoint", "Method", "Description"],
        [
            ["/analyze/form/", "GET/POST", "Multimodal evaluation form"],
            ["/results/<id>/", "GET", "Full analysis result"],
            ["/results/<id>/pdf/", "GET", "Technical analysis report"],
            ["/results/<id>/pitch/pdf/", "GET", "Visual pitch deck"],
            ["/results/<id>/video/generate/", "POST", "Starts video generation"],
            ["/results/<id>/video/progress/<job_id>/", "GET", "Video progress"],
            ["/models/", "GET/POST", "Model management and training"],
            ["/investors/", "GET", "Investor dashboard"],
        ],
    )

    doc.add_heading("6. Usage guide", level=1)
    _add_numbered_list(
        doc,
        [
            "Access the New Pitch page.",
            "Fill in the startup and sector information.",
            "Submit multimodal data (text, document, audio, video, YouTube).",
            "Run the evaluation with the local engine or GPT.",
            "Review the result with score, categories, and recommendations.",
            "Optionally generate an AI video in the desired mode.",
            "Generate the technical PDF report.",
            "Generate a pitch PDF with automatic context-based design.",
            "Generate a pitch PDF with manual premium design (chosen template).",
        ],
    )

    doc.add_heading("7. Operation, testing, and troubleshooting", level=1)
    _add_bullet_list(
        doc,
        [
            "python3 manage.py check",
            "python3 manage.py test",
            "python3 docs/generate_engineering_pdf.py",
            "python3 docs/generate_engineering_docx.py",
        ],
    )
    doc.add_paragraph("Minimum functional checklist:")
    _add_bullet_list(
        doc,
        [
            "Valid multimodal submission.",
            "Score and categories present in the result.",
            "AI video with progress bar.",
            "Pitch PDF with both design modes.",
            "Technical report download.",
        ],
    )
    doc.add_paragraph("Quick troubleshooting:")
    _add_bullet_list(
        doc,
        [
            "D-ID failure: check key, credits, and the image HTTPS URL.",
            "GPT failure: check OPENAI_API_KEY and local fallback.",
            "PDF/DOCX failure: check dependencies and write permissions.",
        ],
    )

    doc.add_heading("8. Documentation publishing", level=1)
    doc.add_paragraph(
        "The project's operational delivery includes sending the updated documentation to the Discord webhook "
        "in PDF and DOCX formats."
    )

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_docx()
    print(path)
