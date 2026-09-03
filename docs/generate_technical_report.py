"""
Gera o relatório técnico académico do projecto StartupScan em formato .docx.

Execução:
    python docs/generate_relatorio_academico.py
"""

import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
DOCS_DIR = os.path.join(ROOT_DIR, "docs")
OUTPUT_PATH = os.path.join(DOCS_DIR, "Relatorio_Tecnico_StartupScan.docx")

# ---------------------------------------------------------------------------
# Metadados académicos
# ---------------------------------------------------------------------------
AUTHOR = "Andre Hangalo"
INSTITUTION = "Universidade Mandume Ya Ndemofayoh"
COURSE = "Especialização em Desenvolvimento de Software"
SUPERVISOR = "Edson Livongue"
ACADEMIC_YEAR = "2026"
PROJECT_TITLE = "StartupScan"
REPORT_SUBTITLE = (
    "Plataforma Inteligente de Avaliação de Startups com Inteligência Artificial"
)
GENERATION_DATE = datetime.now().strftime("%d de %B de %Y").replace(
    "January", "Janeiro").replace("February", "Fevereiro").replace(
    "March", "Março").replace("April", "Abril").replace(
    "May", "Maio").replace("June", "Junho").replace(
    "July", "Julho").replace("August", "Agosto").replace(
    "September", "Setembro").replace("October", "Outubro").replace(
    "November", "Novembro").replace("December", "Dezembro")


# ---------------------------------------------------------------------------
# Helpers de formatação
# ---------------------------------------------------------------------------

def _set_font(run, name="Times New Roman", size=12, bold=False,
              italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def _para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12,
          bold=False, italic=False, space_before=0, space_after=6,
          font="Times New Roman"):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _set_font(run, name=font, size=size, bold=bold, italic=italic)
    return p


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def _bullet(doc, items, size=12):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(item)
        _set_font(run, size=size)


def _numbered(doc, items, size=12):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(item)
        _set_font(run, size=size)


def _table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light List Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for para in cells[i].paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = Inches(col_widths[i])
    return table


def _page_break(doc):
    doc.add_page_break()


def _add_page_numbers(doc):
    """Adiciona numeração de página no rodapé."""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _set_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.0):
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


# ---------------------------------------------------------------------------
# Secções do relatório
# ---------------------------------------------------------------------------

def _capa(doc):
    _para(doc, INSTITUTION, align=WD_ALIGN_PARAGRAPH.CENTER,
          size=14, bold=True, space_after=4)
    _para(doc, COURSE, align=WD_ALIGN_PARAGRAPH.CENTER,
          size=12, space_after=60)

    _para(doc, PROJECT_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER,
          size=20, bold=True, space_after=8)
    _para(doc, REPORT_SUBTITLE, align=WD_ALIGN_PARAGRAPH.CENTER,
          size=14, italic=True, space_after=60)

    _para(doc, f"Autor: {AUTHOR}", align=WD_ALIGN_PARAGRAPH.CENTER,
          size=12, space_after=4)
    _para(doc, f"Orientador: {SUPERVISOR}", align=WD_ALIGN_PARAGRAPH.CENTER,
          size=12, space_after=4)
    _para(doc, f"Ano Lectivo: {ACADEMIC_YEAR}", align=WD_ALIGN_PARAGRAPH.CENTER,
          size=12, space_after=4)
    _para(doc, GENERATION_DATE, align=WD_ALIGN_PARAGRAPH.CENTER,
          size=12, space_after=0)
    _page_break(doc)


def _resumo(doc):
    _heading(doc, "Resumo", level=1)
    _para(doc, (
        "O presente relatório descreve o desenvolvimento do projecto StartupScan, uma plataforma "
        "web full-stack de avaliação inteligente de startups, concebida e implementada no âmbito "
        "do estágio curricular do curso de Especialização em Desenvolvimento de Software da "
        f"{INSTITUTION}. "
        "O sistema tem como objectivo principal automatizar o ciclo completo de análise de pitches "
        "de startups, desde a recepção de entradas multimodais — texto, documentos, áudio, vídeo "
        "e links YouTube — até à produção de artefactos de comunicação prontos para utilização "
        "por empreendedores, analistas e investidores."
    ))
    _para(doc, (
        "A plataforma integra um pipeline de aprendizagem automática baseado em scikit-learn e "
        "XGBoost para geração de um score de sucesso de 0 a 10, com explicabilidade por categorias. "
        "Complementarmente, oferece integração com a API OpenAI GPT como motor alternativo, "
        "geração de relatório técnico em PDF, construção de pitch deck visual em PDF com múltiplos "
        "templates, e geração de vídeo explicativo narrado por inteligência artificial, com suporte "
        "a API D-ID e fallback local via moviepy e TTS."
    ))
    _para(doc, (
        "Do ponto de vista técnico, o sistema foi desenvolvido com Django 6 e Django REST Framework, "
        "utiliza SQLite em desenvolvimento e PostgreSQL em produção, e recorre a Redis e Celery "
        "para processamento assíncrono de tarefas de longa duração. O controlo de acesso é "
        "implementado através de um modelo de papéis (RBAC) com cinco níveis: administrador, "
        "analista, empreendedor, investidor e público geral."
    ))
    _para(doc, (
        "Palavras-chave: inteligência artificial, avaliação de startups, aprendizagem automática, "
        "Django, processamento multimodal, pitch analysis."
    ), italic=True, space_after=0)
    _page_break(doc)


def _indice(doc):
    _heading(doc, "Índice", level=1)
    entries = [
        ("1.", "Introdução", "5"),
        ("2.", "Contexto e Objetivos", "6"),
        ("  2.1.", "Problema", "6"),
        ("  2.2.", "Objetivos do Sistema", "6"),
        ("  2.3.", "Público-Alvo", "7"),
        ("3.", "Arquitectura Técnica", "8"),
        ("  3.1.", "Stack Tecnológico", "8"),
        ("  3.2.", "Arquitectura de Componentes", "9"),
        ("  3.3.", "Fluxo de Dados", "9"),
        ("  3.4.", "Processamento Assíncrono", "10"),
        ("4.", "Modelo de Dados", "11"),
        ("  4.1.", "PitchAnalysis", "11"),
        ("  4.2.", "UserProfile", "12"),
        ("  4.3.", "IdeaPitchSubmission", "12"),
        ("  4.4.", "InvestorConnectionInterest", "13"),
        ("  4.5.", "IdeaPublicFeedback", "13"),
        ("5.", "Funcionalidades Implementadas", "14"),
        ("  5.1.", "Avaliação Multimodal de Pitch", "14"),
        ("  5.2.", "Relatório Técnico PDF", "15"),
        ("  5.3.", "Pitch Deck PDF", "15"),
        ("  5.4.", "Vídeo Explicativo com IA", "16"),
        ("  5.5.", "Construtor de Ideias", "16"),
        ("  5.6.", "Processamento em Lote", "17"),
        ("  5.7.", "Gestão de Modelos de ML", "17"),
        ("  5.8.", "Conexões Investidor-Startup", "17"),
        ("6.", "Controlo de Acesso e Papéis", "18"),
        ("7.", "Referência de Endpoints", "19"),
        ("8.", "Pipeline de Aprendizagem Automática", "21"),
        ("9.", "Instalação e Configuração", "22"),
        ("  9.1.", "Pré-requisitos", "22"),
        ("  9.2.", "Setup Local", "22"),
        ("  9.3.", "Variáveis de Ambiente", "23"),
        ("  9.4.", "Execução com Docker Compose", "25"),
        ("  9.5.", "Deploy na Render", "25"),
        ("10.", "Testes e Validação", "26"),
        ("11.", "Conclusão", "27"),
        ("12.", "Referências", "28"),
    ]
    tbl = doc.add_table(rows=0, cols=3)
    tbl.style = "Table Grid"
    for num, title, page in entries:
        row = tbl.add_row().cells
        row[0].text = num
        row[1].text = title
        row[2].text = page
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
        row[0].width = Inches(0.5)
        row[1].width = Inches(5.0)
        row[2].width = Inches(0.5)
    _page_break(doc)


def _introducao(doc):
    _heading(doc, "1. Introdução", level=1)
    _para(doc, (
        "O ecossistema de startups em Angola e em África de forma geral tem crescido de forma "
        "significativa nos últimos anos, impulsionado pela expansão da conectividade, pelo "
        "surgimento de aceleradoras e pela crescente disponibilidade de capital de risco. "
        "Neste contexto, a capacidade de avaliar oportunidades de investimento de forma "
        "rápida, consistente e fundamentada torna-se um factor crítico para todos os "
        "intervenientes do ecossistema — empreendedores, analistas e investidores."
    ))
    _para(doc, (
        "O projecto StartupScan foi desenvolvido no âmbito do estágio curricular do curso de "
        f"Especialização em Desenvolvimento de Software da {INSTITUTION}, com o objectivo de "
        "responder a esta necessidade através de uma plataforma web que automatiza e padroniza "
        "o processo de avaliação de pitches de startups utilizando inteligência artificial."
    ))
    _para(doc, (
        "O presente relatório técnico documenta a totalidade do trabalho realizado: desde a "
        "definição dos requisitos e a arquitectura do sistema, passando pela implementação de "
        "cada funcionalidade, até aos procedimentos de instalação, configuração, testes e "
        "deployment. O documento está organizado de forma a servir simultaneamente como "
        "referência técnica para desenvolvedores e como relatório académico para avaliação "
        "do trabalho efectuado."
    ))
    _para(doc, (
        "O relatório está estruturado em doze capítulos. Os capítulos 2 e 3 contextualizam o "
        "problema e descrevem a arquitectura do sistema. Os capítulos 4 a 8 detalham o modelo "
        "de dados, as funcionalidades, o controlo de acesso, os endpoints e o pipeline de "
        "aprendizagem automática. Os capítulos 9 e 10 cobrem a instalação e os testes. "
        "O capítulo 11 apresenta as conclusões e o capítulo 12 lista as referências bibliográficas."
    ))


def _contexto_objetivos(doc):
    _heading(doc, "2. Contexto e Objetivos", level=1)

    _heading(doc, "2.1. Problema", level=2)
    _para(doc, (
        "A avaliação de pitches de startups é um processo tradicionalmente manual, demorado e "
        "subjectivo. Um analista experiente consegue avaliar um número limitado de pitches por "
        "semana, e a qualidade do feedback varia significativamente entre avaliadores. "
        "Esta ineficiência cria gargalos em aceleradoras e fundos de investimento, atrasa "
        "o processo de financiamento e priva os empreendedores de feedback estruturado e "
        "accionável em tempo útil."
    ))
    _para(doc, (
        "Adicionalmente, os empreendedores, particularmente os que se encontram em fases iniciais, "
        "carecem frequentemente de ferramentas que os ajudem a estruturar e a comunicar as suas "
        "ideias de forma eficaz para potenciais investidores. A ausência de artefactos de "
        "comunicação de qualidade — pitch decks, relatórios e vídeos de apresentação — "
        "constitui uma barreira ao acesso a financiamento."
    ))

    _heading(doc, "2.2. Objetivos do Sistema", level=2)
    _para(doc, "O StartupScan foi concebido para atingir os seguintes objectivos:")

    _para(doc, "Objectivos de negócio:", bold=True, space_after=2)
    _bullet(doc, [
        "Reduzir o tempo de análise de oportunidades de startup de dias para segundos.",
        "Padronizar a qualidade e a estrutura do feedback para empreendedores.",
        "Fornecer artefactos de comunicação prontos a usar (PDF, vídeo) sem esforço adicional.",
        "Criar um canal estruturado de conexão entre empreendedores e investidores.",
    ])

    _para(doc, "Objectivos técnicos:", bold=True, space_after=2)
    _bullet(doc, [
        "Implementar um pipeline robusto de ingestão de dados multimodais.",
        "Desenvolver e integrar um modelo de aprendizagem automática interpretável.",
        "Construir um sistema de geração automática de artefactos (PDF, vídeo).",
        "Garantir escalabilidade através de processamento assíncrono com Celery.",
        "Disponibilizar uma API RESTful para integração com sistemas externos.",
    ])

    _heading(doc, "2.3. Público-Alvo", level=2)
    _table(doc,
        ["Perfil", "Papel no sistema", "Funcionalidades principais"],
        [
            ["Empreendedor", "empreendedor",
             "Submissão de pitch, construtor de ideias, pitch deck PDF, conexões"],
            ["Analista", "analista",
             "Dashboard analítico, processamento em lote, gestão de modelos ML"],
            ["Investidor", "investidor",
             "Dashboard de deal flow, expressão de interesse, hub de conexões"],
            ["Público geral", "publico_geral",
             "Navegação de ideias públicas, feedback com classificação por estrelas"],
            ["Administrador", "admin",
             "Gestão total: utilizadores, modelos, configuração do sistema"],
        ],
        col_widths=[1.2, 1.2, 3.6],
    )


def _arquitectura(doc):
    _heading(doc, "3. Arquitectura Técnica", level=1)

    _heading(doc, "3.1. Stack Tecnológico", level=2)
    _para(doc, (
        "O sistema foi construído sobre um conjunto de tecnologias maduras e amplamente "
        "adoptadas na indústria, seleccionadas com base em critérios de estabilidade, "
        "ecosistema de suporte e adequação ao problema."
    ))
    _table(doc,
        ["Camada", "Tecnologia", "Versão", "Função"],
        [
            ["Framework web", "Django", "6.0.3", "MVC, ORM, autenticação, admin"],
            ["API REST", "Django REST Framework", "3.16.1", "Serialização, endpoints JSON"],
            ["Servidor", "Gunicorn + WhiteNoise", "23.0 / 6.11", "WSGI, ficheiros estáticos"],
            ["Base de dados (dev)", "SQLite", "3.x", "Persistência local"],
            ["Base de dados (prod)", "PostgreSQL", "15", "Persistência em produção"],
            ["Cache / Filas", "Redis + Celery", "—", "Jobs assíncronos"],
            ["ML / IA", "scikit-learn, XGBoost", "1.4.2 / 2.0.3", "Pipeline de scoring"],
            ["LLM", "OpenAI SDK", "1.23.2", "GPT como motor alternativo"],
            ["Vídeo", "moviepy, OpenCV, D-ID", "2.1.2 / 4.9", "Geração de vídeo IA"],
            ["Áudio / TTS", "Whisper, edge-tts, gTTS", "—", "Transcrição e síntese de voz"],
            ["PDF / DOCX", "ReportLab, pypdf, python-docx", "4.4 / 6.9 / 1.2", "Geração de documentos"],
            ["Frontend", "Django Templates + Bootstrap + Chart.js", "—", "Interface web"],
        ],
        col_widths=[1.5, 1.8, 1.0, 2.2],
    )

    _heading(doc, "3.2. Arquitectura de Componentes", level=2)
    _para(doc, (
        "O sistema segue uma arquitectura em camadas com separação clara de responsabilidades. "
        "Cada camada comunica apenas com a camada adjacente, favorecendo a manutenibilidade "
        "e a testabilidade do código."
    ))
    _table(doc,
        ["Camada", "Módulos", "Responsabilidade"],
        [
            ["Apresentação", "templates/, static/",
             "Interface web, formulários, dashboards, polling de progresso"],
            ["Controlo", "views.py, urls.py",
             "Orquestração de pedidos, controlo de acesso, respostas HTTP"],
            ["Negócio", "services/*, modeling.py",
             "Lógica de análise, geração de artefactos, treino de modelos"],
            ["Dados", "models.py, serializers.py",
             "Abstracção da base de dados, validação, serialização"],
            ["Infra-estrutura", "tasks.py, celery.py, settings.py",
             "Filas assíncronas, configuração, variáveis de ambiente"],
        ],
        col_widths=[1.4, 1.8, 3.3],
    )

    _heading(doc, "3.3. Fluxo de Dados", level=2)
    _para(doc, "O fluxo principal de uma avaliação de pitch segue os seguintes passos:")
    _numbered(doc, [
        "O utilizador submete o pitch através do formulário web ou da API REST.",
        "O módulo pitch_input.py extrai e normaliza o conteúdo de todos os formatos "
        "(texto, PDF, DOCX, áudio via Whisper, vídeo via extracção de áudio).",
        "O motor de feature engineering transforma o texto em vectores TF-IDF e "
        "normaliza os dados financeiros.",
        "O modelo de ML activo (ou GPT, se configurado) gera o score e o relatório "
        "estruturado com categorias e recomendações.",
        "O resultado é persistido na tabela PitchAnalysis da base de dados.",
        "O utilizador é redirecionado para a página de resultados, onde pode descarregar "
        "o relatório PDF, o pitch deck ou iniciar a geração do vídeo.",
    ])

    _heading(doc, "3.4. Processamento Assíncrono", level=2)
    _para(doc, (
        "Tarefas de longa duração — treino de modelos de ML e geração de vídeo — são "
        "executadas de forma assíncrona para não bloquear o servidor web. O padrão "
        "implementado é o seguinte:"
    ))
    _numbered(doc, [
        "O cliente efectua um POST ao endpoint de início de tarefa.",
        "O servidor cria um identificador único de job (job_id) e inicia a tarefa "
        "em background via Celery.",
        "O servidor retorna imediatamente o job_id ao cliente.",
        "O cliente efectua polling periódico ao endpoint de progresso, recebendo "
        "actualizações de estado em JSON.",
        "Quando o estado é 'completed', o cliente apresenta o resultado ao utilizador.",
    ])
    _para(doc, (
        "Este padrão é aplicado tanto ao treino de modelos (/model/retrain/ → "
        "/models/training/progress/<job_id>/) como à geração de vídeo "
        "(/video/generate/ → /video/progress/<job_id>/)."
    ))


def _modelos_dados(doc):
    _heading(doc, "4. Modelo de Dados", level=1)
    _para(doc, (
        "O esquema de dados do StartupScan é composto por cinco modelos principais, "
        "todos geridos pelo ORM do Django e versionados através do sistema de migrações. "
        "A seguir descreve-se cada modelo com os seus campos e relações."
    ))

    _heading(doc, "4.1. PitchAnalysis", level=2)
    _para(doc, (
        "Registo central de cada avaliação de startup. Agrega todos os dados de entrada, "
        "os resultados da análise e os metadados operacionais."
    ))
    _table(doc,
        ["Campo", "Tipo", "Descrição"],
        [
            ["user", "FK User (nullable)", "Utilizador que submeteu a análise"],
            ["startup_name", "CharField", "Nome da startup"],
            ["industry", "CharField", "Sector (tech, health, finance, education, ecommerce, other)"],
            ["contact_email", "EmailField", "Email de contacto"],
            ["text", "TextField", "Texto do pitch"],
            ["audio_file", "FileField", "Ficheiro de áudio enviado"],
            ["video_file", "FileField", "Ficheiro de vídeo enviado"],
            ["document_file", "FileField", "Documento enviado (PDF, DOCX, etc.)"],
            ["presenter_face_image_file", "FileField", "Fotografia do apresentador"],
            ["youtube_url", "URLField", "Link YouTube do pitch"],
            ["revenue", "DecimalField", "Receita em AOA"],
            ["growth_rate", "FloatField", "Taxa de crescimento (%)"],
            ["profit_margin", "FloatField", "Margem de lucro (%)"],
            ["burn_rate", "DecimalField", "Consumo mensal de capital"],
            ["success_score", "FloatField", "Score final de 0 a 10"],
            ["confidence", "FloatField", "Confiança da previsão (%)"],
            ["report", "JSONField", "Relatório estruturado completo"],
            ["status", "CharField", "pending / processing / completed / failed"],
            ["model_version", "CharField", "Versão do modelo de ML utilizado"],
            ["processing_time", "FloatField", "Tempo de processamento em segundos"],
            ["ip_address", "GenericIPAddressField", "IP do cliente"],
            ["created_at", "DateTimeField", "Data/hora de criação"],
            ["updated_at", "DateTimeField", "Data/hora da última actualização"],
        ],
        col_widths=[1.8, 1.5, 3.2],
    )

    _heading(doc, "4.2. UserProfile", level=2)
    _para(doc, (
        "Extensão do modelo de utilizador padrão do Django para suporte ao sistema de papéis. "
        "Criado automaticamente no registo de cada novo utilizador."
    ))
    _table(doc,
        ["Campo", "Tipo", "Descrição"],
        [
            ["user", "OneToOneField User", "Referência ao utilizador Django"],
            ["role", "CharField", "Um dos cinco papéis: empreendedor, investidor, analista, publico_geral, admin"],
            ["created_at", "DateTimeField", "Data/hora de criação"],
            ["updated_at", "DateTimeField", "Data/hora da última actualização"],
        ],
        col_widths=[1.8, 1.8, 2.9],
    )

    _heading(doc, "4.3. IdeaPitchSubmission", level=2)
    _para(doc, (
        "Armazena ideias de negócio submetidas pelo construtor de ideias. Cada registo representa "
        "uma ideia em estado draft ou com pitch gerado."
    ))
    _table(doc,
        ["Campo", "Tipo", "Descrição"],
        [
            ["user", "FK User", "Utilizador criador"],
            ["startup_name", "CharField", "Nome da startup"],
            ["one_liner", "CharField", "Pitch de uma frase"],
            ["problem", "TextField", "Problema que a startup resolve"],
            ["solution", "TextField", "Solução proposta"],
            ["target_customer", "TextField", "Perfil do cliente-alvo"],
            ["market_size", "TextField", "Dimensão do mercado endereçável"],
            ["business_model", "TextField", "Modelo de negócio"],
            ["competitive_advantage", "TextField", "Diferencial competitivo"],
            ["traction", "TextField", "Tracção e validações actuais"],
            ["team", "TextField", "Composição e experiência da equipa"],
            ["funding_goal", "CharField", "Objectivo de financiamento"],
            ["use_of_funds", "TextField", "Alocação prevista do financiamento"],
            ["model_source", "CharField", "Motor de geração: local / gpt"],
            ["status", "CharField", "draft / generated"],
            ["generated_pitch", "JSONField", "Conteúdo do pitch gerado"],
        ],
        col_widths=[1.8, 1.2, 3.5],
    )

    _heading(doc, "4.4. InvestorConnectionInterest", level=2)
    _para(doc, (
        "Regista o interesse de um investidor numa análise de startup. Suporta o ciclo "
        "completo de comunicação entre investidor e empreendedor."
    ))
    _table(doc,
        ["Campo", "Tipo", "Descrição"],
        [
            ["analysis", "FK PitchAnalysis", "Análise de interesse"],
            ["investor", "FK User", "Utilizador investidor"],
            ["entrepreneur", "FK User (nullable)", "Utilizador empreendedor destinatário"],
            ["status", "CharField", "pending / reviewing / connected / rejected / withdrawn"],
            ["investor_message", "TextField", "Mensagem inicial do investidor"],
            ["entrepreneur_reply", "TextField", "Resposta do empreendedor"],
            ["created_at", "DateTimeField", "Data/hora de criação do interesse"],
            ["responded_at", "DateTimeField", "Data/hora da resposta do empreendedor"],
        ],
        col_widths=[1.8, 1.5, 3.2],
    )

    _heading(doc, "4.5. IdeaPublicFeedback", level=2)
    _para(doc, (
        "Avaliação da comunidade sobre ideias tornadas públicas no marketplace de ideias. "
        "Cada utilizador pode submeter apenas um feedback por ideia."
    ))
    _table(doc,
        ["Campo", "Tipo", "Descrição"],
        [
            ["submission", "FK IdeaPitchSubmission", "Ideia avaliada"],
            ["author", "FK User", "Utilizador que avaliou"],
            ["stars", "IntegerField", "Classificação de 1 a 5 estrelas"],
            ["endorsed", "BooleanField", "Indica se o utilizador endossou a ideia"],
            ["comment", "TextField", "Comentário qualitativo"],
            ["created_at", "DateTimeField", "Data/hora da avaliação"],
        ],
        col_widths=[1.8, 1.2, 3.5],
    )


def _funcionalidades(doc):
    _heading(doc, "5. Funcionalidades Implementadas", level=1)

    _heading(doc, "5.1. Avaliação Multimodal de Pitch", level=2)
    _para(doc, (
        "A funcionalidade nuclear do sistema permite que o utilizador submeta informação "
        "sobre a sua startup em qualquer combinação dos seguintes formatos:"
    ))
    _bullet(doc, [
        "Texto livre digitado directamente no formulário.",
        "Documento: .txt, .md, .csv, .pdf, .docx.",
        "Ficheiro de áudio (WAV, MP3, OGG) — transcrito via OpenAI Whisper.",
        "Ficheiro de vídeo — o áudio é extraído e transcrito.",
        "Link YouTube — o áudio é extraído e transcrito.",
        "Dados financeiros: receita (AOA), taxa de crescimento, margem de lucro, burn rate.",
    ])
    _para(doc, (
        "O módulo pitch_input.py consolida todas as entradas num único bloco de texto "
        "normalizado, que é depois processado pelo pipeline de ML. O resultado da análise "
        "inclui: score de 0 a 10, percentagem de confiança, resumo executivo, pontos fortes, "
        "pontos fracos, recomendações práticas e avaliação em oito categorias — Clareza da "
        "Proposta, Proposta de Valor, Inovação, Viabilidade, Escalabilidade, Mercado-Alvo, "
        "Equipa e Sustentabilidade."
    ))

    _heading(doc, "5.2. Relatório Técnico PDF", level=2)
    _para(doc, (
        "Para cada análise concluída, o sistema gera automaticamente um relatório técnico "
        "em formato PDF via ReportLab. O documento inclui: dados da startup submetidos, "
        "score e confiança com visualização gráfica, avaliação detalhada por categoria, "
        "análise financeira, recomendações práticas e metadados do modelo utilizado "
        "(versão, tempo de processamento)."
    ))

    _heading(doc, "5.3. Pitch Deck PDF", level=2)
    _para(doc, (
        "O sistema gera um pitch deck visual em formato PDF, estruturado como slides "
        "prontos para apresentação a investidores. Cada página corresponde a um slide, "
        "com capa executiva, narrativa estruturada e conclusão. "
        "São suportados dois modos:"
    ))
    _bullet(doc, [
        "Design automático por contexto (recomendado): o sistema selecciona o template "
        "com base na indústria e nos dados da startup.",
        "Design premium manual: o utilizador escolhe entre seis templates — Orbit, Grid, "
        "Wave, Diagonal, Aurora e Ribbon.",
    ])

    _heading(doc, "5.4. Vídeo Explicativo com IA", level=2)
    _para(doc, (
        "A partir do resultado da análise, o sistema gera um vídeo narrado por inteligência "
        "artificial com duração entre 1 e 3 minutos. São suportados três modos de geração:"
    ))
    _bullet(doc, [
        "auto: tenta a API D-ID (apresentador realista) e cai para vídeo local em caso de falha.",
        "did_only: utiliza exclusivamente a API D-ID.",
        "local_only: geração local com moviepy e TTS, sem dependências externas pagas.",
    ])
    _para(doc, (
        "O processo é assíncrono, com progresso actualizado em tempo real via polling. "
        "O sistema suporta detecção automática do género do apresentador a partir de uma "
        "fotografia, utilizando o deepface, para seleccionar a voz TTS adequada."
    ))

    _heading(doc, "5.5. Construtor de Ideias", level=2)
    _para(doc, (
        "O construtor de ideias permite ao empreendedor estruturar uma ideia de negócio "
        "através de um formulário guiado com campos para problema, solução, cliente-alvo, "
        "dimensão de mercado, modelo de negócio, vantagem competitiva, tracção, equipa, "
        "objectivo de financiamento e uso dos fundos. A partir destes dados, o sistema "
        "gera automaticamente um pitch narrativo (via modelo local ou GPT), que pode ser "
        "exportado como PDF ou publicado no marketplace público para receber feedback "
        "da comunidade."
    ))

    _heading(doc, "5.6. Processamento em Lote", level=2)
    _para(doc, (
        "Analistas e administradores podem submeter um ficheiro CSV contendo múltiplos "
        "pitches para avaliação em lote. O processamento ocorre de forma assíncrona, "
        "com polling de progresso disponível. Os resultados são disponibilizados para "
        "download num ficheiro CSV consolidado."
    ))

    _heading(doc, "5.7. Gestão de Modelos de ML", level=2)
    _para(doc, (
        "O painel de gestão de modelos, acessível a analistas e administradores, permite:"
    ))
    _bullet(doc, [
        "Importação de datasets externos (CSV de pitches e financeiros).",
        "Treino e retreino do modelo com progresso em tempo real.",
        "Activação do modelo a utilizar nas análises.",
        "Edição de metadados do modelo (nome de exibição, descrição).",
        "Remoção de modelos obsoletos.",
    ])

    _heading(doc, "5.8. Conexões Investidor-Startup", level=2)
    _para(doc, (
        "A plataforma implementa um canal estruturado de comunicação entre investidores "
        "e empreendedores. O investidor expressa interesse numa análise com uma mensagem "
        "personalizada. O empreendedor recebe a notificação no hub de conexões e pode "
        "aceitar, rejeitar ou responder. O estado da conexão percorre o ciclo: "
        "pending → reviewing → connected / rejected."
    ))


def _controlo_acesso(doc):
    _heading(doc, "6. Controlo de Acesso e Papéis", level=1)
    _para(doc, (
        "O sistema implementa controlo de acesso baseado em papéis (RBAC — Role-Based "
        "Access Control). Cada utilizador possui um UserProfile com um papel atribuído "
        "no momento do registo. O acesso a cada funcionalidade é verificado pelo mixin "
        "RoleRequiredMixin nas views. Administradores correspondem a superutilizadores "
        "Django, criados via manage.py createsuperuser."
    ))
    _table(doc,
        ["Papel", "Dashboard", "Pitch", "Modelos ML", "Investidor", "Ideias", "Conexões", "Admin"],
        [
            ["admin", "✓", "✓", "✓", "✓", "✓", "✓", "✓"],
            ["analista", "✓", "✓", "✓", "—", "✓", "—", "—"],
            ["empreendedor", "✓", "✓", "—", "—", "✓", "✓", "—"],
            ["investidor", "—", "—", "—", "✓", "✓", "✓", "—"],
            ["publico_geral", "—", "—", "—", "—", "✓ (ver)", "—", "—"],
        ],
        col_widths=[1.2, 0.8, 0.7, 0.9, 0.9, 0.7, 0.9, 0.7],
    )
    _para(doc, (
        "A função role_home_url() determina a página inicial de cada utilizador após "
        "autenticação. As views sensíveis utilizam LoginRequiredMixin em combinação com "
        "RoleRequiredMixin para garantir dupla verificação."
    ))


def _endpoints(doc):
    _heading(doc, "7. Referência de Endpoints", level=1)
    _para(doc, (
        "A tabela seguinte lista todos os endpoints disponíveis na plataforma, "
        "organizados por domínio funcional."
    ))

    _heading(doc, "Autenticação", level=2)
    _table(doc,
        ["Método", "Endpoint", "Descrição"],
        [
            ["GET/POST", "/login/", "Página e processamento de login"],
            ["GET", "/logout/", "Encerrar sessão"],
            ["GET/POST", "/register/", "Página e processamento de registo"],
            ["POST", "/set-language/", "Alteração do idioma da interface"],
        ],
        col_widths=[0.8, 2.0, 3.7],
    )

    _heading(doc, "Análise de Pitch", level=2)
    _table(doc,
        ["Método", "Endpoint", "Descrição"],
        [
            ["GET/POST", "/analyze/form/", "Formulário e submissão de pitch"],
            ["POST", "/analyze/", "Endpoint API REST de análise"],
            ["GET", "/results/<id>/", "Página de resultados"],
            ["GET", "/results/<id>/pdf/", "Relatório técnico PDF"],
            ["GET", "/results/<id>/pitch/pdf/", "Pitch deck PDF"],
        ],
        col_widths=[0.8, 2.5, 3.2],
    )

    _heading(doc, "Vídeo Explicativo", level=2)
    _table(doc,
        ["Método", "Endpoint", "Descrição"],
        [
            ["POST", "/results/<id>/video/generate/", "Iniciar geração de vídeo"],
            ["POST", "/results/<id>/video/detect-gender/", "Detectar género do apresentador"],
            ["GET", "/results/<id>/video/progress/<job_id>/", "Polling do progresso do vídeo"],
        ],
        col_widths=[0.8, 3.0, 2.7],
    )

    _heading(doc, "Processamento em Lote", level=2)
    _table(doc,
        ["Método", "Endpoint", "Descrição"],
        [
            ["POST", "/batch/analyze/", "Submeter CSV para análise em lote"],
            ["GET", "/batch/status/<batch_id>/", "Estado do processamento em lote"],
            ["GET", "/batch/results/<batch_id>/", "Descarregar resultados em CSV"],
        ],
        col_widths=[0.8, 2.5, 3.2],
    )

    _heading(doc, "Gestão de Modelos", level=2)
    _table(doc,
        ["Método", "Endpoint", "Descrição"],
        [
            ["GET", "/models/", "Painel de gestão de modelos"],
            ["POST", "/model/retrain/", "Iniciar treino de modelo"],
            ["GET", "/models/training/progress/<job_id>/", "Progresso do treino"],
            ["GET", "/training/status/<task_id>/", "Estado de task Celery"],
        ],
        col_widths=[0.8, 2.7, 3.0],
    )

    _heading(doc, "Construtor de Ideias e Conexões", level=2)
    _table(doc,
        ["Método", "Endpoint", "Descrição"],
        [
            ["GET/POST", "/pitch/builder/", "Formulário e submissão de ideia"],
            ["GET", "/pitch/builder/<id>/", "Detalhe e edição de ideia"],
            ["GET", "/pitch/builder/<id>/pdf/", "Exportar ideia como PDF"],
            ["GET", "/ideas/", "Marketplace de ideias públicas"],
            ["GET", "/ideas/<id>/", "Detalhe de ideia pública"],
            ["POST", "/ideas/<id>/feedback/", "Submeter feedback a ideia"],
            ["POST", "/investors/interest/<id>/", "Expressar interesse numa startup"],
            ["GET", "/connections/", "Hub de conexões"],
            ["POST", "/connections/<id>/update/", "Actualizar estado de conexão"],
        ],
        col_widths=[0.8, 2.5, 3.2],
    )


def _pipeline_ml(doc):
    _heading(doc, "8. Pipeline de Aprendizagem Automática", level=1)
    _para(doc, (
        "O motor de scoring do StartupScan é um pipeline de aprendizagem automática "
        "supervisionada, treinado com dados históricos de pitches de startups. "
        "A seguir descreve-se cada etapa do pipeline."
    ))

    _heading(doc, "8.1. Pré-processamento e Feature Engineering", level=2)
    _bullet(doc, [
        "Texto do pitch: vectorização TF-IDF com n-gramas de 1 a 2 tokens.",
        "Dados financeiros: normalização com StandardScaler (receita, crescimento, "
        "margem de lucro, burn rate).",
        "Feature de saúde financeira: métrica composta calculada a partir do crescimento, "
        "margem e receita.",
        "Augmentação de dados: factor de 60x com jitter gaussiano para enriquecer "
        "conjuntos de dados escassos.",
    ])

    _heading(doc, "8.2. Modelo Ensemble", level=2)
    _para(doc, (
        "O modelo utiliza um ensemble de três estimadores base com votação suave:"
    ))
    _bullet(doc, [
        "Random Forest Classifier (scikit-learn).",
        "Gradient Boosting Classifier (scikit-learn).",
        "Extra Trees Classifier (scikit-learn).",
    ])
    _para(doc, (
        "Como alternativa, o XGBoost Classifier pode ser activado para datasets maiores. "
        "A validação cruzada utiliza KFold com 5 partições para estimativa de desempenho "
        "não enviesada."
    ))

    _heading(doc, "8.3. Saída do Modelo", level=2)
    _bullet(doc, [
        "Score de sucesso: valor contínuo de 0 a 10.",
        "Confiança: percentagem derivada da dispersão entre os estimadores do ensemble.",
        "Categorias: oito dimensões avaliadas individualmente com base nos features "
        "mais relevantes para cada dimensão.",
        "Recomendações: texto interpretável gerado com base nas categorias com menor score.",
    ])

    _heading(doc, "8.4. Integração GPT", level=2)
    _para(doc, (
        "Quando a variável OPENAI_API_KEY está configurada, o sistema pode utilizar "
        "GPT como motor de análise alternativo. A função analyze_with_gpt() envia o "
        "pitch consolidado para a API OpenAI e processa a resposta no mesmo formato "
        "estruturado que o modelo local. Se a API não estiver disponível ou falhar, "
        "o sistema cai automaticamente para o modelo local sem intervenção do utilizador."
    ))

    _heading(doc, "8.5. Gestão do Ciclo de Vida do Modelo", level=2)
    _para(doc, (
        "O modelo activo é gerido pelo módulo model_registry.py, que mantém um ficheiro "
        "de metadados JSON com o nome e o caminho do modelo activo. Novos modelos treinados "
        "são persistidos como ficheiros .pkl via joblib. A activação de um novo modelo "
        "actualiza imediatamente o registo, afectando todas as análises subsequentes."
    ))


def _instalacao(doc):
    _heading(doc, "9. Instalação e Configuração", level=1)

    _heading(doc, "9.1. Pré-requisitos", level=2)
    _table(doc,
        ["Software", "Versão mínima", "Obrigatório", "Nota"],
        [
            ["Python", "3.10", "Sim", "Recomendado 3.11 ou 3.12"],
            ["pip", "23.x", "Sim", "Incluído com Python"],
            ["Git", "2.x", "Sim", "Para clonar o repositório"],
            ["Redis", "7.x", "Não*", "Necessário para Celery (jobs assíncronos)"],
            ["FFmpeg", "6.x", "Não*", "Necessário para geração de vídeo local"],
            ["Docker", "24.x", "Não*", "Para execução containerizada"],
        ],
        col_widths=[1.2, 1.0, 1.0, 3.3],
    )
    _para(doc, (
        "* Opcional para o funcionamento básico (avaliação, PDF). "
        "Redis e FFmpeg são necessários para vídeo local e processamento assíncrono."
    ), italic=True)

    _heading(doc, "9.2. Setup Local", level=2)
    _numbered(doc, [
        "Clonar o repositório: git clone https://github.com/rickdeu/startupscan-backend.git",
        "Criar ambiente virtual: python -m venv .venv",
        "Activar o ambiente virtual: source .venv/bin/activate (Linux/macOS) "
        "ou .venv\\Scripts\\activate (Windows)",
        "Instalar dependências: pip install -r requirements.txt",
        "Criar ficheiro .env com as variáveis de ambiente necessárias (ver 9.3).",
        "Aplicar migrações: python manage.py migrate",
        "Criar conta de administrador: python manage.py createsuperuser",
        "Recolher ficheiros estáticos: python manage.py collectstatic --noinput",
        "(Opcional) Treinar modelo inicial: python manage.py train_model "
        "--model-output ai_models/pitch_model.pkl",
        "Iniciar o servidor: python manage.py runserver 0.0.0.0:8000",
    ])

    _heading(doc, "9.3. Variáveis de Ambiente", level=2)
    _para(doc, "Variáveis obrigatórias:", bold=True, space_after=2)
    _table(doc,
        ["Variável", "Descrição", "Exemplo"],
        [
            ["SECRET_KEY", "Chave secreta Django (obrigatória em produção)",
             "django-insecure-..."],
            ["DJANGO_DEBUG", "Modo debug: 1 para dev, 0 para prod", "1"],
        ],
        col_widths=[1.8, 2.7, 2.0],
    )
    _para(doc, "Variáveis de base de dados (para PostgreSQL):", bold=True, space_after=2)
    _table(doc,
        ["Variável", "Descrição", "Exemplo"],
        [
            ["DATABASE_URL", "URL de conexão PostgreSQL completa",
             "postgres://user:pass@host:5432/db"],
            ["POSTGRES_USER", "Utilizador PostgreSQL", "startupscan"],
            ["POSTGRES_PASSWORD", "Password PostgreSQL", "password123"],
            ["POSTGRES_DB", "Nome da base de dados", "startupscan"],
            ["POSTGRES_HOST", "Host do servidor PostgreSQL", "localhost"],
            ["POSTGRES_PORT", "Porta PostgreSQL", "5432"],
        ],
        col_widths=[1.8, 2.2, 2.5],
    )
    _para(doc, "Variáveis de APIs externas:", bold=True, space_after=2)
    _table(doc,
        ["Variável", "Efeito se ausente"],
        [
            ["OPENAI_API_KEY", "Sistema usa exclusivamente o modelo local de ML"],
            ["OPENAI_MODEL", "Usa gpt-4.1-mini como modelo GPT"],
            ["DID_API_KEY", "Geração de vídeo usa modo local (moviepy + TTS)"],
            ["DID_API_BASE_URL", "Usa https://api.d-id.com como endpoint"],
            ["EDGE_TTS_VOICE_PT_AO", "Usa voz padrão do edge-tts em português"],
            ["WHISPER_MODEL", "Usa modelo 'base' para transcrição de áudio"],
        ],
        col_widths=[2.2, 4.3],
    )

    _heading(doc, "9.4. Execução com Docker Compose", level=2)
    _para(doc, (
        "O ficheiro docker-compose.yml disponível na raiz do projecto levanta o stack "
        "completo com um único comando:"
    ))
    _bullet(doc, [
        "web: aplicação Django via Gunicorn na porta 8000.",
        "db: PostgreSQL 15 na porta 5432.",
        "redis: Redis 7 na porta 6379.",
        "celery-worker: processamento assíncrono de tarefas.",
        "celery-beat: agendamento de tarefas periódicas.",
    ])
    _para(doc, "Comando de arranque: docker-compose up -d")
    _para(doc, "Comando para parar: docker-compose down")
    _para(doc, "Remover todos os dados: docker-compose down -v")

    _heading(doc, "9.5. Deploy na Render", level=2)
    _para(doc, (
        "O projecto está configurado para deploy contínuo na plataforma Render.com "
        "através do ficheiro render.yaml e do workflow GitHub Actions em "
        ".github/workflows/deploy-render-main.yml. "
        "O deploy é despoletado automaticamente por cada push para a branch main, "
        "após configuração do secret RENDER_DEPLOY_HOOK_URL no repositório GitHub."
    ))


def _testes(doc):
    _heading(doc, "10. Testes e Validação", level=1)

    _heading(doc, "10.1. Testes Automáticos", level=2)
    _para(doc, (
        "A suite de testes automáticos está implementada em startupscan_api/tests.py "
        "e cobre os principais fluxos da aplicação. Para executar os testes:"
    ))
    _bullet(doc, [
        "python manage.py test — executa toda a suite de testes.",
        "python manage.py check — verifica a configuração do sistema Django.",
    ])

    _heading(doc, "10.2. Checklist de Validação Funcional", level=2)
    _para(doc, (
        "Após a instalação ou após alterações significativas ao código, "
        "recomenda-se a validação dos seguintes fluxos:"
    ))
    _numbered(doc, [
        "Registo e autenticação com cada um dos cinco papéis.",
        "Submissão de pitch com texto simples e verificação do score gerado.",
        "Submissão de pitch com ficheiro PDF e verificação da extracção de texto.",
        "Download do relatório técnico PDF.",
        "Geração de pitch deck PDF em modo automático.",
        "Geração de pitch deck PDF em modo premium (pelo menos um template).",
        "Geração de vídeo em modo local_only.",
        "Polling de progresso de geração de vídeo (verificar actualizações).",
        "Treino de modelo via painel e activação do novo modelo.",
        "Polling de progresso de treino de modelo.",
        "Submissão de ideia no construtor e exportação como PDF.",
        "Publicação de ideia e submissão de feedback pela conta de público geral.",
        "Fluxo completo de conexão: expressão de interesse → resposta do empreendedor.",
    ])

    _heading(doc, "10.3. Troubleshooting", level=2)
    _table(doc,
        ["Problema", "Causa provável", "Solução"],
        [
            ["SECRET_KEY não definida",
             "DJANGO_DEBUG=0 sem SECRET_KEY no .env",
             "Adicionar SECRET_KEY ao .env ou definir DJANGO_DEBUG=1"],
            ["Vídeo D-ID falha",
             "API key inválida, sem créditos, ou imagem não acessível via HTTPS",
             "Verificar DID_API_KEY e créditos; testar com modo local_only"],
            ["PDF não gera",
             "ReportLab não instalado ou MEDIA_ROOT sem permissão de escrita",
             "pip install reportlab; verificar permissões do directório media/"],
            ["GPT não é utilizado",
             "OPENAI_API_KEY ausente ou inválida",
             "Definir OPENAI_API_KEY no .env; sistema usa modelo local como fallback"],
            ["Celery não processa jobs",
             "Redis não está a correr",
             "Iniciar Redis (redis-server) e worker Celery"],
            ["Overlay de submissão persiste",
             "Cache do browser",
             "Limpar cache do browser (Ctrl+Shift+R)"],
        ],
        col_widths=[1.5, 2.0, 3.0],
    )


def _conclusao(doc):
    _heading(doc, "11. Conclusão", level=1)
    _para(doc, (
        "O projecto StartupScan demonstrou ser tecnicamente viável e funcionalmente completo "
        "no contexto do estágio curricular. A plataforma implementa com sucesso o ciclo "
        "completo de avaliação de startups — da ingestão multimodal de dados à produção "
        "de artefactos de comunicação — e integra um conjunto diversificado de tecnologias "
        "de ponta em aprendizagem automática, processamento de linguagem natural, "
        "síntese de vídeo e geração de documentos."
    ))
    _para(doc, (
        "Do ponto de vista académico, o projecto permitiu aplicar em contexto real um "
        "conjunto abrangente de conhecimentos adquiridos ao longo do curso de especialização: "
        "arquitectura de software web, desenvolvimento de APIs RESTful, implementação de "
        "pipelines de machine learning, processamento assíncrono de tarefas, gestão de "
        "bases de dados relacionais, e práticas de DevOps com Docker e CI/CD."
    ))
    _para(doc, (
        "Os principais desafios técnicos enfrentados durante o desenvolvimento incluíram: "
        "a normalização de entradas multimodais com formatos e qualidades variáveis, "
        "a calibração do modelo de ML com datasets de dimensão limitada (resolvida com "
        "augmentação de dados), a integração de APIs externas com comportamentos "
        "assíncronos (D-ID, OpenAI), e a implementação de um sistema de polling "
        "eficiente para jobs de longa duração sem degradar a experiência do utilizador."
    ))
    _para(doc, (
        "Como trabalho futuro, identificam-se as seguintes oportunidades de evolução: "
        "integração de modelos de linguagem de grande escala (LLM) mais recentes para "
        "análise semântica mais profunda; implementação de um sistema de recomendação "
        "para ligar empreendedores a investidores com base em preferências de sector; "
        "dashboard de métricas e analytics para administradores; e suporte a múltiplos "
        "idiomas na geração de artefactos."
    ))
    _para(doc, (
        "Em suma, o StartupScan representa uma contribuição concreta para a modernização "
        "do ecossistema de avaliação de startups, com potencial de aplicação em aceleradoras, "
        "fundos de investimento e programas de empreendedorismo no contexto angolano e africano."
    ))


def _referencias(doc):
    _heading(doc, "12. Referências", level=1)
    refs = [
        ("Django Software Foundation", "Django Web Framework", "2024",
         "https://www.djangoproject.com/"),
        ("Tom Christie", "Django REST Framework", "2024",
         "https://www.django-rest-framework.org/"),
        ("Pedregosa, F. et al.", "Scikit-learn: Machine Learning in Python. "
         "Journal of Machine Learning Research, 12, 2825–2830.", "2011", "—"),
        ("Chen, T. & Guestrin, C.", "XGBoost: A Scalable Tree Boosting System. "
         "Proceedings of the 22nd ACM SIGKDD.", "2016", "—"),
        ("OpenAI", "GPT-4 Technical Report", "2024", "https://openai.com/research/gpt-4"),
        ("D-ID", "D-ID API Documentation", "2024", "https://docs.d-id.com/"),
        ("Radford, A. et al.", "Robust Speech Recognition via Large-Scale Weak Supervision "
         "(Whisper)", "2022", "https://openai.com/research/whisper"),
        ("Zulko", "MoviePy: Video Editing with Python", "2024",
         "https://zulko.github.io/moviepy/"),
        ("ReportLab", "ReportLab PDF Library", "2024", "https://www.reportlab.com/"),
        ("Docker Inc.", "Docker Documentation", "2024", "https://docs.docker.com/"),
        ("Render", "Render Documentation", "2024", "https://render.com/docs"),
        ("Celery Project", "Celery: Distributed Task Queue", "2024",
         "https://docs.celeryq.dev/"),
    ]
    for i, (authors, title, year, url) in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        r1 = p.add_run(f"[{i}] {authors}. ")
        _set_font(r1, size=11)
        r2 = p.add_run(f"{title}. ")
        _set_font(r2, size=11, italic=True)
        r3 = p.add_run(f"{year}.")
        _set_font(r3, size=11)
        if url != "—":
            r4 = p.add_run(f" Disponível em: {url}")
            _set_font(r4, size=11)


# ---------------------------------------------------------------------------
# Builder principal
# ---------------------------------------------------------------------------

def build():
    os.makedirs(DOCS_DIR, exist_ok=True)
    doc = Document()
    _set_margins(doc, top=1.1, bottom=1.0, left=1.25, right=1.0)
    _add_page_numbers(doc)

    _capa(doc)
    _resumo(doc)
    _indice(doc)
    _introducao(doc)
    _page_break(doc)
    _contexto_objetivos(doc)
    _page_break(doc)
    _arquitectura(doc)
    _page_break(doc)
    _modelos_dados(doc)
    _page_break(doc)
    _funcionalidades(doc)
    _page_break(doc)
    _controlo_acesso(doc)
    _page_break(doc)
    _endpoints(doc)
    _page_break(doc)
    _pipeline_ml(doc)
    _page_break(doc)
    _instalacao(doc)
    _page_break(doc)
    _testes(doc)
    _page_break(doc)
    _conclusao(doc)
    _page_break(doc)
    _referencias(doc)

    doc.save(OUTPUT_PATH)
    print(f"Relatório gerado em: {OUTPUT_PATH}")
    return OUTPUT_PATH


if __name__ == "__main__":
    build()
